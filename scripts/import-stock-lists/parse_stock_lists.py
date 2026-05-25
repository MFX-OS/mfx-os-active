"""
parse_stock_lists.py
--------------------
Parse the STD Flexible Film List and STD Label Stock List .docx files into
MATS-shaped entries and write the result to ./new_mats.json.

Schema produced (matches the MATS literal in public/js/core.js):
  { v: vendor, s: spec, d: description, m: msi cost, mk: msi marked-up,
    liner?, adhesive? }   # liner+adhesive only on label-stock entries

Usage:
  python3 parse_stock_lists.py

Inputs:
  source-docs/STD-Flexible-Film-List.docx
  source-docs/STD-Label-Stock-List.docx

Output:
  new_mats.json  (consumed by build_mats_patch.py)
"""
from docx import Document
import json
import os
import sys
from collections import Counter

HERE = os.path.dirname(os.path.abspath(__file__))
SRC_DIR = os.path.join(HERE, 'source-docs')
OUT = os.path.join(HERE, 'new_mats.json')

# Specs already present in core.js MATS — keep this list in sync if the
# hard-coded MATS literal changes. (Used only for visibility; the patch
# script also dedupes against the actual file contents at apply time.)
EXISTING_SPECS = {
    '539SHP', '162WP', '317HSB', '329RPP', '371WOPP', '378WPET', '338SHP',
    '401WPET', '480SPET', '211HSF', '362WSF', '143SW', '264SW', '531SW',
    '3515W', '557SW', '523CBOPP', '524WPET', '5525F', '444 STML', '464 GSCL',
    '133 SG', '238 WBOPP', '281SG', '393SG', '171WBOPP', '450WP', '515NKL',
    '295 GCF', '408 CMOP', '412 SHP', '391SG',
}


def infer_vendor(vendor_ref):
    """Map a free-text vendor reference (e.g. 'Avery Spec# 14470') to a vendor key."""
    if not vendor_ref:
        return 'STD'
    s = vendor_ref.lower()
    if 'avery' in s:
        return 'Avery'
    if 'green bay' in s or 'greenbay' in s:
        return 'Greenbay'
    if 'klockner' in s:
        return 'KLOCKNER'
    if 'qspac' in s:
        return 'QSPAC'
    if 'achem' in s:
        return 'ACHEM'
    if 'pac-master' in s or 'pacmaster' in s:
        return 'PAC-MASTER'
    if 'nobelus' in s:
        return 'Nobelus'
    if 'klaser' in s:
        return 'Klaser'
    if 'spinnaker' in s:
        return 'Spinnaker'
    if 'multi-plastics' in s or 'multiplastics' in s:
        return 'MULTI-PLASTICS'
    return 'STD'


def parse_films(path):
    """Films doc: 2-row-per-SKU pattern.
       Row A: [spec, description, description]
       Row B: [spec, 'Avery Spec# N' or similar vendor ref, ...]"""
    d = Document(path)
    rows = []
    for t in d.tables:
        for row in t.rows:
            rows.append([c.text.strip() for c in row.cells])
    entries = []
    i = 1  # skip header
    while i < len(rows):
        if i + 1 < len(rows) and rows[i][0] == rows[i + 1][0]:
            entries.append({
                's': rows[i][0],
                'd': rows[i][1],
                'v': infer_vendor(rows[i + 1][1]),
                'm': None, 'mk': None,
                '_vref': rows[i + 1][1],
            })
            i += 2
        else:
            spec = rows[i][0]
            if spec and spec != 'Our Spec':
                entries.append({
                    's': spec, 'd': rows[i][1], 'v': 'STD',
                    'm': None, 'mk': None, '_vref': '',
                })
            i += 1
    return entries


def parse_labels(path):
    """Labels doc: 2-row-per-SKU pattern.
       Row A: [spec, description, description, description]
       Row B: [spec, vendor-ref, liner, adhesive]"""
    d = Document(path)
    rows = []
    for t in d.tables:
        for row in t.rows:
            rows.append([c.text.strip() for c in row.cells])
    entries = []
    i = 2  # skip 2 header rows
    while i < len(rows):
        if i + 1 < len(rows) and rows[i][0] == rows[i + 1][0]:
            vref = rows[i + 1][1] if len(rows[i + 1]) > 1 else ''
            liner = rows[i + 1][2] if len(rows[i + 1]) > 2 else ''
            adhesive = rows[i + 1][3] if len(rows[i + 1]) > 3 else ''
            entries.append({
                's': rows[i][0],
                'd': rows[i][1],
                'v': infer_vendor(vref),
                'm': None, 'mk': None,
                'liner': liner or None,
                'adhesive': adhesive or None,
                '_vref': vref,
            })
            i += 2
        else:
            i += 1
    return entries


def main():
    films_path = os.path.join(SRC_DIR, 'STD-Flexible-Film-List.docx')
    labels_path = os.path.join(SRC_DIR, 'STD-Label-Stock-List.docx')

    if not os.path.exists(films_path) or not os.path.exists(labels_path):
        print(f'ERROR: expected source docs in {SRC_DIR}', file=sys.stderr)
        sys.exit(1)

    films = parse_films(films_path)
    labels = parse_labels(labels_path)

    all_new = []
    seen = set(EXISTING_SPECS)
    dup_existing = dup_within = 0
    for e in films + labels:
        spec = e['s'].strip()
        if spec in seen:
            if spec in EXISTING_SPECS:
                dup_existing += 1
            else:
                dup_within += 1
            continue
        seen.add(spec)
        all_new.append(e)

    vendor_breakdown = Counter(e['v'] for e in all_new)

    print('=== PARSE RESULTS ===')
    print(f'Films parsed:        {len(films)}')
    print(f'Labels parsed:       {len(labels)}')
    print(f'Total raw:           {len(films) + len(labels)}')
    print(f'Dups w/ existing:    {dup_existing}')
    print(f'Dups within docs:    {dup_within}')
    print(f'NET NEW:             {len(all_new)}')
    print()
    print('=== VENDOR BREAKDOWN ===')
    for v, n in vendor_breakdown.most_common():
        print(f'  {v}: {n}')

    with open(OUT, 'w') as f:
        json.dump(all_new, f, indent=2)
    print(f'\nWrote {len(all_new)} entries to {OUT}')


if __name__ == '__main__':
    main()
